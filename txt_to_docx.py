#!/usr/bin/env python3
import os
import argparse
import tempfile
from docx import Document
import gc

READ_BUFFER_SIZE = 2 * 1024 * 1024       # 2 MB per read
MAX_SPLIT_SIZE = 240 * 1024 * 1024       # 480 MB per DOCX file

def split_txt_file(txt_path, tmp_dir):
    """
    Split a large TXT file into 480MB chunk files in tmp_dir.
    Returns a list of chunk file paths.
    """
    buffers_per_chunk = MAX_SPLIT_SIZE // READ_BUFFER_SIZE
    chunk_paths = []

    with open(txt_path, 'r', encoding='utf-8') as f:
        chunk_index = 1
        while True:
            
            chunk_path = os.path.join(tmp_dir, f"{chunk_index}.txt")
            print(f"Creating split file: {chunk_path}")
            
            chunk_paths.append(chunk_path)
            with open(chunk_path, 'w', encoding='utf-8') as out:
                for _ in range(buffers_per_chunk):
                    data = f.read(READ_BUFFER_SIZE)
                    if not data:
                        return chunk_paths  # EOF reached
                    out.write(data)
            
            chunk_index += 1

def txt_to_docx(txt_path, docx_path):
    """Convert a single TXT file to DOCX (memory-safe, reads in buffers)."""
    doc = Document()
    paragraph = doc.add_paragraph()
    with open(txt_path, 'r', encoding='utf-8') as f:
        while True:
            data = f.read(READ_BUFFER_SIZE)
            if not data:
                break
            paragraph.add_run(data)
    doc.save(docx_path)
    print(f"Saved {docx_path}")
    input("Press Enter to continue...")  # Pause for user to see progress
    
# Memory leak mitigation: force garbage collection after each conversion
def txt_to_docx_with_gc(txt_path, docx_path):
    txt_to_docx(txt_path, docx_path)
    gc.collect()
    

def convert_large_txt(txt_path, output_name):
    """Convert TXT to DOCX. Split into multiple DOCX if >480MB."""
    
    
    file_size = os.path.getsize(txt_path)
    print(f"Input TXT size: {file_size / (1024*1024):.2f} MB")

    if file_size <= MAX_SPLIT_SIZE:
        # Small file: single DOCX
        txt_to_docx(txt_path, f"{output_name}.docx")
    else:
        # Large file: split and convert
        os.makedirs(output_name, exist_ok=True)
        with tempfile.TemporaryDirectory() as tmp_dir:
            
            print(f"Splitting TXT into 480MB chunks in temporary dir: {tmp_dir}")
            chunk_files = split_txt_file(txt_path, tmp_dir)
            
            for i, chunk_file in enumerate(chunk_files, start=1):
                                
                docx_path = os.path.join(output_name, f"{i}.docx")
                txt_to_docx_with_gc(chunk_file, docx_path)                
                
            print(f"All chunks converted. DOCX files are in {output_name}/")

def main():
    parser = argparse.ArgumentParser(description="Convert TXT to DOCX (split 480MB each if needed)")
    parser.add_argument("txt_file", help="Input TXT file path")
    parser.add_argument("--output-name", required=True, help="Output file name (for small files) or directory name (for large files)")
    args = parser.parse_args()

    convert_large_txt(args.txt_file, args.output_name)

if __name__ == "__main__":
    main()
