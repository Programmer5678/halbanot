#!/usr/bin/env python3
import os

import tempfile
from docx import Document
import gc

READ_BUFFER_SIZE = 2 * 1024 * 1024       # 2 MB per read
MAX_SPLIT_SIZE = 15 * 1024 * 1024       # 15 MB per DOCX file

def split_txt_file(txt_path, tmp_dir):
    """
    Split a large TXT file into 15MB chunk files in tmp_dir.
    Returns a list of chunk file paths.
    """
    buffers_per_chunk = MAX_SPLIT_SIZE // READ_BUFFER_SIZE
    chunk_paths = []

    with open(txt_path, 'r', encoding='utf-8') as f:
        chunk_index = 1
        while True:
            
            chunk_path = os.path.join(tmp_dir, f"{chunk_index}.txt")
            print(f"Creating split file: {chunk_path}")
            
            chunk_paths.append(chunk_path)
            with open(chunk_path, 'w', encoding='utf-8') as out:
                for _ in range(buffers_per_chunk):
                    data = f.read(READ_BUFFER_SIZE)
                    if not data:
                        return chunk_paths  # EOF reached
                    out.write(data)
            
            chunk_index += 1

def txt_to_docx_single_file(txt_path, docx_path):
    """Convert a single TXT file to DOCX (memory-safe, reads in buffers)."""
    doc = Document()
    paragraph = doc.add_paragraph()
    with open(txt_path, 'r', encoding='utf-8') as f:
        while True:
            data = f.read(READ_BUFFER_SIZE)
            if not data:
                break
            paragraph.add_run(data)
    doc.save(docx_path)
    print(f"Saved {docx_path}")

    return docx_path
    
# Memory leak mitigation: force garbage collection after each conversion
def txt_to_docx_with_gc(txt_path, docx_path):
    txt_to_docx_single_file(txt_path, docx_path)
    gc.collect()
 
 
 
 
def convert_chunks_to_docx(chunk_files, output_base, max_files_per_dir=40):
    """
    Convert a list of TXT chunk files to DOCX, splitting into subdirectories
    with a maximum number of files per directory.
    
    Args:
        chunk_files (list[str]): List of TXT chunk file paths.
        output_base (str): Base directory to store DOCX subfolders.
        max_files_per_dir (int): Max number of DOCX files per subdirectory.
    """
    for i, chunk_file in enumerate(chunk_files, start=1):
        # Determine which subdir this file belongs to
        dir_index = (i - 1) // max_files_per_dir
        dir_name = os.path.join(output_base, f"{dir_index:04d}")
        os.makedirs(dir_name, exist_ok=True)

        # DOCX filename uses global count
        docx_path = os.path.join(dir_name, f"{i:04d}.docx")
        txt_to_docx_with_gc(chunk_file, docx_path)


def to_docx_path(output_path):
    return output_path if output_path.endswith(".docx") else f"{output_path}.docx"

def txt_to_docxs(txt_path, output_path, max_files_per_dir=40) -> str:
    """Convert TXT to DOCX. Split into multiple DOCX if >15MB,
    organizing DOCX files into directories with max_files_per_dir each.

    Return full output path
    """
    
    file_size = os.path.getsize(txt_path)
    print(f"Input TXT size: {file_size / (1024*1024):.2f} MB")

    if file_size <= MAX_SPLIT_SIZE:
        # Small file: single DOCX

        result = to_docx_path(output_path)
        txt_to_docx_single_file(txt_path, result)
    else:
        result = output_path
        # Large file: split and convert
        txt_to_dir_of_docxs(max_files_per_dir, result, txt_path)

    return result


def txt_to_dir_of_docxs(max_files_per_dir, output_path, txt_path) -> str:
    os.makedirs(output_path, exist_ok=True)
    with tempfile.TemporaryDirectory() as tmp_dir:
        print(f"Splitting TXT into chunks in temporary dir: {tmp_dir}")
        chunk_files = split_txt_file(txt_path, tmp_dir)

        # Use modular function
        convert_chunks_to_docx(chunk_files, output_path, max_files_per_dir)

        print(f"All chunks converted. DOCX files are in subfolders of {output_path}/")



